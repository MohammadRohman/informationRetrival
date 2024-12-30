import os
from django.core.wsgi import get_wsgi_application
from django.urls import path
from django.conf import settings
from django.conf.urls.static import static
from django.shortcuts import render
from django.http import HttpResponse
from django.core.files.storage import FileSystemStorage
from django.forms import Form, FileField, CharField
import nltk
from nltk.corpus import stopwords
from Sastrawi.Stemmer.StemmerFactory import StemmerFactory
from PyPDF2 import PdfReader
from docx import Document
import string
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.decomposition import TruncatedSVD
from sklearn.metrics.pairwise import cosine_similarity
import time
from collections import Counter

# Django settings
DEBUG = True
SECRET_KEY = 'django-insecure-your-random-string'
BASE_DIR = os.path.dirname(os.path.abspath(__file__))

ALLOWED_HOSTS = []

MEDIA_URL = '/media/'
MEDIA_ROOT = os.path.join(BASE_DIR, 'media', 'documents')
STATIC_URL = '/static/'
STATIC_ROOT = os.path.join(BASE_DIR, 'static')

# Create required directories
os.makedirs(MEDIA_ROOT, exist_ok=True)
os.makedirs(STATIC_ROOT, exist_ok=True)

# Configure Django
settings.configure(
    DEBUG=DEBUG,
    SECRET_KEY=SECRET_KEY,
    ROOT_URLCONF=__name__,
    ALLOWED_HOSTS=ALLOWED_HOSTS,
    MIDDLEWARE=[
        'django.middleware.security.SecurityMiddleware',
        'django.contrib.sessions.middleware.SessionMiddleware',
        'django.middleware.common.CommonMiddleware',
        'django.middleware.csrf.CsrfViewMiddleware',
        'django.contrib.auth.middleware.AuthenticationMiddleware',
        'django.contrib.messages.middleware.MessageMiddleware',
    ],
    INSTALLED_APPS=[
        'django.contrib.auth',
        'django.contrib.contenttypes',
        'django.contrib.sessions',
        'django.contrib.messages',
        'django.contrib.staticfiles',
    ],
    TEMPLATES=[{
        'BACKEND': 'django.template.backends.django.DjangoTemplates',
        'DIRS': [os.path.join(BASE_DIR, 'templates')],
        'APP_DIRS': True,
        'OPTIONS': {
            'context_processors': [
                'django.template.context_processors.debug',
                'django.template.context_processors.request',
                'django.contrib.auth.context_processors.auth',
                'django.contrib.messages.context_processors.messages',
            ],
        },
    }],
    DATABASES={
        'default': {
            'ENGINE': 'django.db.backends.sqlite3',
            'NAME': os.path.join(BASE_DIR, 'db.sqlite3'),
        }
    },
    MEDIA_URL=MEDIA_URL,
    MEDIA_ROOT=MEDIA_ROOT,
    STATIC_URL=STATIC_URL,
    STATIC_ROOT=STATIC_ROOT,
    BASE_DIR=BASE_DIR
)

import django
django.setup()

# Text processing setup
nltk.download('stopwords')
nltk.download('punkt')
stop_words = set(stopwords.words('indonesian'))
factory = StemmerFactory()
stemmer = factory.create_stemmer()

def preprocess_text(text):
    text = text.lower()
    tokens = nltk.word_tokenize(text.translate(str.maketrans('', '', string.punctuation)))
    stemmed_tokens = [stemmer.stem(word) for word in tokens if word not in stop_words]
    word_count = Counter(stemmed_tokens)
    return ' '.join(stemmed_tokens), len(stemmed_tokens), word_count

def read_file(file):
    content = ""
    try:
        if isinstance(file, str):  # If file is a path
            if file.endswith('.txt'):
                with open(file, 'r', encoding='utf-8') as f:
                    content = f.read()
            elif file.endswith('.pdf'):
                from PyPDF2 import PdfReader
                reader = PdfReader(file)
                content = "".join([page.extract_text() for page in reader.pages])
            elif file.endswith('.docx'):
                from docx import Document
                doc = Document(file)
                content = "\n".join([p.text for p in doc.paragraphs])
            else:
                raise ValueError("Unsupported file type.")
        else:  # If file is a file object
            if file.name.endswith('.txt'):
                content = file.read().decode('utf-8')
            elif file.name.endswith('.pdf'):
                from PyPDF2 import PdfReader
                reader = PdfReader(file)
                content = "".join([page.extract_text() for page in reader.pages])
            elif file.name.endswith('.docx'):
                from docx import Document
                doc = Document(file)
                content = "\n".join([p.text for p in doc.paragraphs])
            else:
                raise ValueError("Unsupported file type.")
    except Exception as e:
        content = f"Error reading file: {e}"
    return content

# Forms
class UploadForm(Form):
    file = FileField()

class SearchForm(Form):
    query = CharField()

# Views
def get_available_files():
    files = [f for f in os.listdir(MEDIA_ROOT) if os.path.isfile(os.path.join(MEDIA_ROOT, f))]
    return files

def view_document(request, doc_name):
    doc_name = os.path.basename(doc_name)
    file_content = os.path.join(MEDIA_ROOT, doc_name)
    # file_path = os.path.join(MEDIA_ROOT, f"{doc_name}.processed")
    count_path = os.path.join(MEDIA_ROOT, f"{doc_name}.count")
    stemming_path = os.path.join(MEDIA_ROOT, f"{doc_name}.stemming")
    document_content = ""
    word_count = 0
    stemming_content = []
    
    # read content file
    document_content = read_file(file_content)
    
    if os.path.exists(count_path):
        with open(count_path, 'r', encoding='utf-8') as f:
            word_count = int(f.read())
    
    if os.path.exists(stemming_path):
        with open(stemming_path, 'r', encoding='utf-8') as f:
            stemming_content = f.readlines()
    
    return render(request, 'document_view.html', {
        'document_content': document_content,
        'doc_name': doc_name,
        'word_count': word_count,
        'stemming_content': stemming_content,
    })

def upload_view(request):
    available_files = get_available_files()

    if request.method == 'POST':
        form = UploadForm(request.POST, request.FILES)
        if form.is_valid():
            file = request.FILES['file']
            fs = FileSystemStorage()
            fs.save(file.name, file)

            content = read_file(file)
            processed_content, word_count, word_count_dict = preprocess_text(content)

            with open(os.path.join(MEDIA_ROOT, f"{file.name}.processed"), 'w', encoding='utf-8') as f:
                f.write(processed_content)

            with open(os.path.join(MEDIA_ROOT, f"{file.name}.count"), 'w', encoding='utf-8') as f:
                f.write(str(word_count))

            with open(os.path.join(MEDIA_ROOT, f"{file.name}.stemming"), 'w', encoding='utf-8') as f:
                for word, count in word_count_dict.items():
                    original_words = [token for token in nltk.word_tokenize(content) if stemmer.stem(token.lower()) == word]
                    original_words_str = ', '.join(original_words)
                    f.write(f"Kata Dasar: {word} ({count} kali) -> Asli: {original_words_str}\n")

            return render(request, 'upload.html', {'form': UploadForm(), 'message': 'File uploaded successfully!', 'available_files': available_files})

    return render(request, 'upload.html', {'form': UploadForm(), 'available_files': available_files})

def search_view(request):
    available_files = get_available_files()

    if request.method == 'POST':
        form = SearchForm(request.POST)
        if form.is_valid():
            query = form.cleaned_data['query']
            documents = {}
            for filename in available_files:
                if filename.endswith('.processed'):
                    with open(os.path.join(MEDIA_ROOT, filename), 'r', encoding='utf-8') as f:
                        documents[filename.replace('.processed', '')] = f.read()

            if documents:
                vectorizer = TfidfVectorizer()
                tfidf_matrix = vectorizer.fit_transform(documents.values())
                svd = TruncatedSVD(n_components=min(100, tfidf_matrix.shape[1]-1), random_state=42)
                lsi_matrix = svd.fit_transform(tfidf_matrix)
                
                query_processed, _, _ = preprocess_text(query)
                query_tfidf = vectorizer.transform([query_processed])
                query_lsi = svd.transform(query_tfidf)
                
                similarities = cosine_similarity(query_lsi, lsi_matrix)[0]
                results = sorted(zip(documents.keys(), similarities), key=lambda x: x[1], reverse=True)
                return render(request, 'search.html', {'form': form, 'results': results, 'available_files': available_files})
    
    return render(request, 'search.html', {'form': SearchForm(), 'available_files': available_files})

# URLs
urlpatterns = [
    path('', upload_view, name='upload'),
    path('search/', search_view, name='search'),
    path('view_document/<str:doc_name>/', view_document, name='view_document'),
] + static(settings.MEDIA_URL, document_root=settings.MEDIA_ROOT)
# Templates
UPLOAD_TEMPLATE = '''
<!DOCTYPE html>
<html>
<head>
    <title>Upload Document</title>
    <link href="https://cdn.jsdelivr.net/npm/tailwindcss@2.2.19/dist/tailwind.min.css" rel="stylesheet">
</head>
<body class="bg-gray-100">
    <div class="container mx-auto px-4 py-8">
        <div class ="max-w-md mx-auto bg-white rounded-xl shadow-md overflow-hidden p-6">
            <h2 class="text-2xl font-bold mb-4">Upload Document</h2>
            {% if message %}<p class="text-green-500 mb-4">{{ message }}</p>{% endif %}
            <form method="post" enctype="multipart/form-data">
                {% csrf_token %}
                <div class="mb-4">{{ form.file }}</div>
                <button type="submit" class="bg-green-500 text-white px-4 py-2 rounded hover:bg-blue-600">Upload</button>
            </form>
            <a href="/search" class="mt-4 inline-block text-blue-500">Go to Search</a>
        </div>
    </div>
</body>
</html>
'''

SEARCH_TEMPLATE = '''
<!DOCTYPE html>
<html>
<head>
    <title>Search Documents</title>
    <link href="https://cdn.jsdelivr.net/npm/tailwindcss@2.2.19/dist/tailwind.min.css" rel="stylesheet">
</head>
<body class="bg-gray-100">
    <div class="container mx-auto px-4 py-8">
        <div class="max-w-md mx-auto bg-white rounded-xl shadow-md overflow-hidden p-6">
            <h2 class="text-2xl font-bold mb-4">Search Documents</h2>
            <form method="post">
                {% csrf_token %}
                <div class="mb-4">{{ form.query }}</div>
                <button type="submit" class="bg-blue-500 text-white px-4 py-2 rounded hover:bg-blue-600">Search</button>
            </form>
            {% if results %}
            <div class="mt-6">
                <h3 class="text-xl font-bold mb-2">Results:</h3>
                <ul>
                {% for doc, similarity in results %}
                    <li class="mb-2">
                        <a href="{% url 'view_document' doc %}" class="flex justify-between text-blue-500 hover:underline">
                            <span>{{ doc }}</span>
                            <span class="text-gray-600">{{ similarity|floatformat:4 }}</span>
                        </a>
                    </li>
                {% endfor %}
                </ul>
            </div>
            {% endif %}
            <a href="/" class="mt-4 inline-block text-blue-500">Upload New Document</a>
        </div>
    </div>
</body>
</html>
'''

DOCUMENT_VIEW_TEMPLATE = '''
<!DOCTYPE html>
<html>
<head>
    <title>Document Content</title>
    <link href="https://cdn.jsdelivr.net/npm/tailwindcss@2.2.19/dist/tailwind.min.css" rel="stylesheet">
</head>
<body class="bg-gray-100">
    <div class="container mx-auto px-12 py-24">
        <div class="max-w-md mx-auto bg-white rounded-xl shadow-md overflow-hidden p-8">
            <h2 class="text-2xl font-bold mb-4">{{ doc_name }} - Document Content</h2>
            
            <div class="mb-4 p-4 border border-gray-300 rounded-lg bg-gray-50">
                <h3 class="text-lg font-semibold">Word Count (Stemmed):</h3>
                <p class="text-gray-700">{{ word_count }}</p>
            </div>

            <div class="mb-4 p-4 border border-gray-300 rounded-lg bg-gray-50">
                 <h3 class="text-lg font-semibold">Stemmed Words:</h3>
                <div class="bg-gray-100 p-4 rounded">
                    <ul>
                        {% for line in stemming_content %}
                        <li class="text-gray-700">{{ line }}</li>
                        {% endfor %}
                    </ul>
                </div>
            </div>

            <div class="mt-4">
                <h3 class="text-lg font-semibold">Document Content:</h3>
                <p class="text-gray-700">{{ document_content }}</p>
            </div>

            <a href="/search" class="mt-4 inline-block text-blue-500">Back to Search</a>
        </div>
    </div>
</body>
</html>
'''

# Create templates directory and files
os.makedirs(os.path.join(BASE_DIR, 'templates'), exist_ok=True)
os.makedirs(MEDIA_ROOT, exist_ok=True)

with open(os.path.join(BASE_DIR, 'templates', 'upload.html'), 'w') as f:
    f.write(UPLOAD_TEMPLATE)
with open(os.path.join(BASE_DIR, 'templates', 'search.html'), 'w') as f:
    f.write(SEARCH_TEMPLATE)
with open(os.path.join(BASE_DIR, 'templates', 'document_view.html'), 'w') as f:
    f.write(DOCUMENT_VIEW_TEMPLATE)

# Run server
if __name__ == '__main__':
    from django.core.management import execute_from_command_line
    execute_from_command_line(['manage.py', 'runserver'])
